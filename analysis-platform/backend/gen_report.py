# -*- coding: utf-8 -*-
"""소통 플랫폼 및 분석 플랫폼 구축 보고서 (.docx) 생성.

서식 규격: Arial(코드만 Consolas), 제목 진남색(1F4E79)→소제목 파랑(2E75B6)→항목 회색(333333).
구성: 헤더박스 / 표지 / 목차 / 본문(H1~H3·표·코드·주석) / 그림캡션 / 푸터.
"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Mm

OUT = Path(r"c:\Users\Administrator\OneDrive\Desktop\소통 플랫폼 및 분석 플랫폼 구축 보고서.docx")

NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x33, 0x33, 0x33)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DIMGREY = RGBColor(0x70, 0x70, 0x70)
HDR_NAVY = "1F4E79"
HDR_BLUE = "2E75B6"
ZEBRA = "F3F6FB"
CODEBG = "F2F2F2"
NOTEBG = "EAF1FB"

doc = Document()


# ── 폰트(한·영 Arial) ──────────────────────────────────────────────
def _font(run, name="Arial", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), name)


def _style_font(style, name="Arial", size=11, bold=False, color=None):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), name)


# 본문/제목 스타일(목차 자동생성을 위해 내장 Heading 스타일 사용)
_style_font(doc.styles["Normal"], "Arial", 11, False, BLACK)
_style_font(doc.styles["Heading 1"], "Arial", 15, True, NAVY)
_style_font(doc.styles["Heading 2"], "Arial", 12, True, BLUE)
_style_font(doc.styles["Heading 3"], "Arial", 10.5, True, GREY)
for hs in ("Heading 1", "Heading 2", "Heading 3"):
    pf = doc.styles[hs].paragraph_format
    pf.space_before = Pt(10 if hs == "Heading 1" else 7)
    pf.space_after = Pt(4)
    pf.keep_with_next = True


# 목차(TOC) 항목 스타일 — 줄 간격 타이트하게.
# 핵심: customStyle 속성을 제거해 Word 내장 "TOC n" 스타일을 '덮어쓰기'로 인식시킨다.
# (custom으로 두면 필드 갱신 시 Word가 자기 내장 TOC 스타일=간격 큼 을 써서 무시됨)
def _toc_style(name, size):
    try:
        st = doc.styles[name]
    except KeyError:
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    el = st.element
    el.set(qn("w:customStyle"), "0")      # 내장 스타일로 취급 → TOC 갱신이 이 정의를 사용
    bo = el.find(qn("w:basedOn"))
    if bo is None:
        bo = OxmlElement("w:basedOn"); el.insert(0, bo)
    bo.set(qn("w:val"), "Normal")
    _style_font(st, "Arial", size, False, BLACK)
    p = st.paragraph_format
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    p.line_spacing = 1.0
    p.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.keep_with_next = False
    # 직접 spacing 강제(twips): before/after 0, 단줄 240
    sp = p.element.get_or_add_pPr().find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing"); p.element.get_or_add_pPr().append(sp)
    sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "240"); sp.set(qn("w:lineRule"), "auto")
    sp.set(qn("w:beforeAutospacing"), "0"); sp.set(qn("w:afterAutospacing"), "0")


for _tn, _ts in (("TOC 1", 11), ("TOC 2", 10.5), ("TOC 3", 10)):
    _toc_style(_tn, _ts)


def H1(t): doc.add_paragraph(t, style="Heading 1")
def H2(t): doc.add_paragraph(t, style="Heading 2")
def H3(t): doc.add_paragraph(t, style="Heading 3")


def P(t, size=11, color=BLACK, bold=False, italic=False, after=4, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Mm(indent)
    r = p.add_run(t)
    _font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def BULLET(t, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t)
    _font(r, size=size, color=BLACK)
    return p


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), fill)
    tcPr.append(sh)


def _cellfont(cell, text, size=9.5, bold=False, color=BLACK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    _font(r, size=size, bold=bold, color=color)


def _left_bar(cell, color):
    """셀 좌측 세로줄(주석/코드 박스용)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "0"); left.set(qn("w:color"), color)
    borders.append(left)
    tcPr.append(borders)


def TABLE(headers, rows, header_fill=HDR_NAVY, widths=None, hsize=9.5, bsize=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        _shade(t.rows[0].cells[i], header_fill)
        _cellfont(t.rows[0].cells[i], h, size=hsize, bold=True, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 1:
                _shade(cells[ci], ZEBRA)
            _cellfont(cells[ci], str(val), size=bsize)
    if widths:
        for col, w in enumerate(widths):
            for r in t.rows:
                r.cells[col].width = Mm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def CODE(text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    _shade(c, CODEBG)
    _left_bar(c, "8FA0B3")
    c.text = ""
    first = True
    for line in text.rstrip("\n").split("\n"):
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else " ")
        _font(r, "Consolas", 9, color=RGBColor(0x22, 0x33, 0x44))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def NOTE(text, title="참고"):
    # WAS 템플릿 스타일: 음영 박스 없이 인라인 '주석' 텍스트
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("주석  "); _font(r, size=9.5, bold=True, color=BLUE)
    body = text if title in ("참고", "목차 안내") else f"{title} — {text}"
    r2 = p.add_run(body); _font(r2, size=9.5, color=GREY)


def CAPTION(n, desc, loc=None):
    # WAS 템플릿 스타일: 인라인 '[그림 N]  설명 (위치: …)'
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"[그림 {n}]  "); _font(r, size=9.5, bold=True, color=NAVY)
    tail = desc + (f" (위치: {loc})" if loc else "")
    r2 = p.add_run(tail); _font(r2, size=9.5, color=GREY)


def field(p, instr, ph="1"):
    """워드 필드(TOC/PAGE) 삽입. ph=갱신 전 임시 표시값(생성 후 Word 자동화로 baked)."""
    r1 = p.add_run(); f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1._r.append(f1)
    r2 = p.add_run(); it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr; r2._r.append(it)
    r3 = p.add_run(); f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "separate"); r3._r.append(f3)
    r4 = p.add_run(ph); _font(r4, size=10, color=DIMGREY)
    r5 = p.add_run(); f5 = OxmlElement("w:fldChar"); f5.set(qn("w:fldCharType"), "end"); r5._r.append(f5)


TODAY = "2026-06-20"
REVDATE = "2026-06-23"
DOCNO = "SOC-BUILD-2026-001"
VER = "1.3"

# ── 페이지 설정 (WAS 템플릿: 머리글 박스 없음 / 바닥글 페이지번호만 유지) ──
sec = doc.sections[0]
sec.top_margin = Mm(22); sec.bottom_margin = Mm(18)
sec.left_margin = Mm(22); sec.right_margin = Mm(22)
# 머리글 없음. 표지·목차 섹션은 바닥글도 비움 → 본문 섹션부터 페이지번호 1로 매김.


def _start_body_section():
    """본문 시작 — 새 섹션(새 페이지) + 페이지 번호를 1부터 재시작 + 본문 전용 푸터(페이지 번호)."""
    new = doc.add_section(WD_SECTION.NEW_PAGE)
    sectPr = new._sectPr
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "1")              # 이 섹션부터 페이지 번호 1로 시작
    sectPr.append(pg)
    new.footer.is_linked_to_previous = False   # 앞(표지·목차)과 분리된 푸터
    fp = new.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field(fp, "PAGE")                           # 페이지 번호 숫자만 표시
    for r in fp.runs:                            # 필드 글꼴 정리
        _font(r, size=9, color=DIMGREY)
    return new


# ── 표지 (WAS 템플릿: 제목 / 부제 / 범위 / 날짜 — 간결, 정보표 없음) ──
def cover():
    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = CENTER
    r = p.add_run("SOC 통합 보안관제 플랫폼"); _font(r, size=28, bold=True, color=NAVY)
    p2 = doc.add_paragraph(); p2.alignment = CENTER
    r2 = p2.add_run("구축 보고서"); _font(r2, size=28, bold=True, color=NAVY)
    doc.add_paragraph()
    ps = doc.add_paragraph(); ps.alignment = CENTER
    rs = ps.add_run("분석플랫폼(SOC Analysis) · 소통플랫폼(SOC Communication) 기반 AI 보안관제")
    _font(rs, size=13, color=BLUE)
    pl = doc.add_paragraph(); pl.alignment = CENTER
    rl = pl.add_run("Splunk 상관분석 · AI 정·오탐 판정 · 협업 티켓팅 · 위협 인텔리전스")
    _font(rl, size=10.5, color=DIMGREY)
    for _ in range(7):
        doc.add_paragraph()
    pd = doc.add_paragraph(); pd.alignment = CENTER
    rd = pd.add_run(TODAY.replace("-", ".")); _font(rd, size=13, bold=True, color=GREY)
    pn = doc.add_paragraph(); pn.alignment = CENTER
    rn = pn.add_run(f"문서번호 {DOCNO}  ·  Ver {VER}  ·  보안관제팀 · 정보보호팀")
    _font(rn, size=9.5, color=DIMGREY)
    doc.add_page_break()


# ── 목차 (WAS 템플릿: '■ 목차') ────────────────────────────────────
def toc():
    p = doc.add_paragraph(); r = p.add_run("■ 목차"); _font(r, size=15, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(6)
    tp = doc.add_paragraph()
    field(tp, 'TOC \\o "1-3" \\h \\z \\u', ph="목차")


cover()
toc()
_start_body_section()   # 여기부터 본문 — 페이지 번호 1로 재시작

# ════════════════════════════════════════════════════════════════════
# 1. 문서 개요
# ════════════════════════════════════════════════════════════════════
H1("1. 문서 개요")
H2("1.1 목적")
P("본 문서는 보안관제센터(SOC) 운영을 위해 구축한 두 개의 웹 플랫폼 — 위협을 탐지·분석·판정하는 "
  "‘분석플랫폼’과, 판정된 정탐 이벤트를 보안관제팀·정보보호팀·웹관리자가 협업으로 처리(티켓팅)하는 "
  "‘소통플랫폼’ — 의 구축 결과를 정리한다. 시스템 구성, 기능 정의서, API 명세서, DB 명세서, "
  "데이터 처리 파이프라인, AI·상관분석 설계, 보안 고려사항을 포함한다.")
H2("1.2 범위")
BULLET("분석플랫폼: Splunk SIEM 상관분석 결과(notable) 수신, AI 기반 정·오탐 자동 판정, 위협 인텔리전스·CVE 보강, 대시보드/지구본 시각화, PDF 보고서, 로그 분석 챗봇")
BULLET("소통플랫폼: 정탐 이벤트 수신·티켓팅(라이프사이클), 협업(댓글·이력·작업), IOC 관리, 사내 메일(janus.com 게이트웨이), 채팅·DM, 대장(원장)·통계")
BULLET("연동: Splunk REST, 메일 게이트웨이(IMAP/SMTP), OpenRouter(LLM), AbuseIPDB/OTX(위협 인텔)")
H2("1.3 용어 정의")
TABLE(["용어", "설명"], [
    ["SIEM", "Security Information and Event Management — 보안 로그 통합·상관분석 시스템(본 구축은 Splunk)"],
    ["상관분석(Correlation)", "다중 보안장비 로그를 규칙으로 엮어 의미 있는 ‘인시던트’를 탐지하는 기법"],
    ["notable", "상관룰이 생성한 탐지 이벤트(soc_notable_json 인덱스에 JSON으로 적재)"],
    ["정탐/오탐", "True Positive / False Positive — 탐지가 실제 공격인지(정탐) 정상인지(오탐)"],
    ["티켓팅", "탐지 이벤트를 사건(티켓)으로 등록해 접수→검토→대응→종결로 처리하는 프로세스"],
    ["IOC", "Indicator of Compromise — 침해지표(악성 IP·도메인·해시 등)"],
    ["MITRE ATT&CK", "공격 전술·기법 분류 체계(예: T1190, T1059.007)"],
    ["SLA", "Service Level Agreement — 우선순위별 처리 기한(P1 1h~P4 72h)"],
], widths=[40, 126])

# ════════════════════════════════════════════════════════════════════
# 2. 시스템 개요
# ════════════════════════════════════════════════════════════════════
H1("2. 시스템 개요")
H2("2.1 배경 및 목표")
P("기존에는 보안장비(WAF·IDS/Snort·방화벽)가 각기 대량의 단일 알림을 쏟아내어 분석가가 일일이 "
  "확인하기 어려웠다. 이를 해결하기 위해 ① Splunk에서 상관분석으로 ‘의미 있는 탐지’만 추려내고, "
  "② 분석플랫폼이 AI로 정·오탐을 자동 판정하여 정탐만 골라, ③ 소통플랫폼에서 팀 간 협업으로 "
  "신속히 대응·종결하는 end-to-end 관제 체계를 구축하는 것을 목표로 한다.")
H2("2.2 플랫폼 구성")
TABLE(["구분", "분석플랫폼", "소통플랫폼"], [
    ["역할", "탐지·분석·정/오탐 판정", "협업 대응·티켓팅·소통"],
    ["포트", "8800", "8810"],
    ["백엔드", "FastAPI (Python)", "FastAPI (Python)"],
    ["프런트엔드", "React (CRA, /frontend/build)", "React (Vite, /frontend/dist)"],
    ["데이터", "Splunk 인덱스 + JSON 캐시(무DB)", "SQLite (SQLAlchemy ORM)"],
    ["AI", "OpenRouter gemini-2.5-flash", "(분석플랫폼 판정 결과 수신)"],
    ["주요 화면", "알림·정오탐·기록·인텔리전스·인사이트·보고서", "대시보드·이벤트·대장·IOC·메일·채팅"],
], widths=[28, 69, 69])
H2("2.3 전체 아키텍처")
P("데이터 흐름은 ‘보안장비 → Splunk 상관분석 → 분석플랫폼(AI 판정) → 소통플랫폼(티켓팅)’ 단방향 "
  "파이프라인이다. 정탐만 소통플랫폼으로 전달되어 인시던트당 한 건의 티켓으로 관리된다.")
CODE("[보안장비]            [SIEM]                  [분석플랫폼:8800]            [소통플랫폼:8810]\n"
     "WAF(modsec)  ┐                                                      \n"
     "IDS(Snort)   ├─▶ Splunk 상관룰 ──▶ soc_notable_json ──▶ 수신·AI 정/오탐 판정 ──▶ 정탐 전달 ──▶ 티켓팅\n"
     "pfSense/감사 ┘     (CORR-0xx)        (notable JSON)        (gemini-2.5-flash)      (dedup)     (라이프사이클)\n"
     "                                                              │                                  │\n"
     "                                      위협인텔(AbuseIPDB/OTX)·CVE 보강            메일/채팅/DM·IOC·대장")
CAPTION(1, "SOC 통합 보안관제 플랫폼 전체 아키텍처 / 데이터 파이프라인", "2.3 전체 아키텍처")
H2("2.4 기술 스택")
TABLE(["계층", "기술"], [
    ["백엔드", "Python 3.13, FastAPI, Uvicorn, SQLAlchemy(소통), requests"],
    ["프런트엔드", "React, axios, react-globe.gl(지구본), Vite/CRA"],
    ["데이터/검색", "Splunk(REST search/jobs, HEC), SQLite"],
    ["AI/LLM", "OpenRouter API(OpenAI 호환), google/gemini-2.5-flash"],
    ["외부 연동", "메일 게이트웨이(IMAP 993/SMTP 25), AbuseIPDB, AlienVault OTX, KrCERT/보안뉴스"],
    ["문서/보고", "reportlab(PDF), python-docx(보고서)"],
], widths=[30, 136])

# ════════════════════════════════════════════════════════════════════
# 3. 시스템 구성 및 연동
# ════════════════════════════════════════════════════════════════════
H1("3. 시스템 구성 및 연동")
H2("3.1 네트워크 / 포트")
TABLE(["구성요소", "주소·포트", "비고"], [
    ["분석플랫폼", "0.0.0.0:8800", "FastAPI, 프런트 정적서빙 포함"],
    ["소통플랫폼", "0.0.0.0:8810", "FastAPI, WebSocket /ws"],
    ["Splunk", "10.0.200.201:8089", "REST(search/jobs), 토큰/계정 인증"],
    ["메일 게이트웨이", "192.168.126.222(→10.0.10.50)", "IMAP 993 / SMTP 25, janus.com"],
    ["OpenRouter", "openrouter.ai (HTTPS)", "LLM 추론(FW 아웃바운드 허용)"],
], widths=[34, 64, 68])
H2("3.2 외부 연동")
TABLE(["연동 대상", "방식", "용도"], [
    ["Splunk", "REST /services/search/jobs, HEC", "상관 notable·원시 alert 조회, 집계"],
    ["메일 게이트웨이", "IMAP4_SSL(993)/SMTP(25)", "사용자별 사서함 송수신·휴지통"],
    ["OpenRouter(LLM)", "POST /v1/chat/completions", "정·오탐 판정·보고서·챗봇·인사이트"],
    ["AbuseIPDB / OTX", "REST API", "출발지 IP 평판(위협 인텔리전스)"],
    ["KrCERT / 보안뉴스", "HTML/RSS 크롤링(캐시)", "위기경보단계·보안권고·보안뉴스"],
], widths=[34, 62, 70])
NOTE("위협 인텔·뉴스 크롤링은 제목·링크만 수집(본문 미복제)하고 30분 캐시·출처표기로 운영한다. "
     "메일 게이트웨이 비밀번호 등 비밀값은 환경변수(.env)로만 주입하며 코드·문서에 포함하지 않는다.", "보안 운영")

# ════════════════════════════════════════════════════════════════════
# 4. 기능 정의서
# ════════════════════════════════════════════════════════════════════
H1("4. 기능 정의서")
H2("4.1 분석플랫폼 기능")
TABLE(["기능 ID", "기능명", "설명"], [
    ["AN-01", "상관탐지 수신", "Splunk soc_notable_json에서 상관룰 탐지(JSON)를 받아 알림 스키마로 매핑"],
    ["AN-02", "AI 정·오탐 자동판정", "신규 그룹만 LLM 판정(캐시 재사용), 페이로드/행위 근거 기반"],
    ["AN-03", "판정 보정(가드레일)", "결정적 공격구문·DDoS 볼륨·상관탐지 신뢰 기반 자동 보정(오탐 방지)"],
    ["AN-04", "중복제거(correlation)", "유형+출발지/자산+(페이로드)+시간 단위로 ‘한 공격=한 건’ 병합"],
    ["AN-05", "탐지 기록(전체 DB뷰)", "기간·검색·필터·페이지네이션으로 전체 탐지 이력 조회"],
    ["AN-06", "위협 인텔리전스", "출발지 IP 평판(AbuseIPDB/OTX) 단건·일괄 조회"],
    ["AN-07", "CVE 연계", "최근 고위험 CVE 목록 + 시그니처-CVE 매핑"],
    ["AN-08", "대시보드·지구본", "금일 위협현황·위기경보·뉴스/권고, 공격 출발지→목적지 호(arc) 시각화"],
    ["AN-09", "인사이트", "기간 트렌드 통계 + LLM 해석, 주간 위협 요약(경영진용)"],
    ["AN-10", "PDF 보고서", "이벤트 분석 결과를 디자인된 PDF로 생성·다운로드"],
    ["AN-11", "로그 분석 챗봇", "수집 로그(24h)를 컨텍스트로 자연어 Q&A"],
    ["AN-12", "정탐 전달", "정탐만 선별해 소통플랫폼으로 전달(중복 전송 방지)"],
], widths=[18, 36, 112])
H2("4.2 소통플랫폼 기능")
TABLE(["기능 ID", "기능명", "설명"], [
    ["CO-01", "정탐 이벤트 수신", "분석플랫폼 /ingest로 정탐 수신, 티켓번호·우선순위·SLA 자동 부여"],
    ["CO-02", "티켓 라이프사이클", "미접수→접수→검토→대응→승인대기→종결(오탐/무시 종결 포함)"],
    ["CO-03", "협업", "댓글·상태이력(감사추적)·담당자 배정·대응 작업(체크리스트)·증적 첨부"],
    ["CO-04", "수동 티켓", "분석플랫폼 외 수동 사건 등록"],
    ["CO-05", "대장(원장)·통계", "전체 처리 대장 조회·CSV 내보내기, 지표/진행현황 대시보드"],
    ["CO-06", "IOC 관리", "침해지표 등록·차단상태 관리, 이벤트에서 자동 추출"],
    ["CO-07", "사내 메일", "janus.com 사서함 송수신·휴지통·검색·첨부파일(IMAP/SMTP 게이트웨이)"],
    ["CO-08", "채팅·DM", "팀 채널 채팅·1:1 다이렉트 메시지(이미지·이벤트카드 공유)"],
    ["CO-09", "단계별 알림", "티켓 단계 전환 시 담당 팀에 메일·텔레그램 통지"],
    ["CO-10", "연락처·부서", "사내 부서 주소록(외부 메일 발송 대상)"],
    ["CO-11", "MITRE 매핑", "공격유형→MITRE 기법 ID 매핑(룰 큐레이션 우선)"],
    ["CO-12", "실시간(WebSocket)", "신규 이벤트·채팅·DM 실시간 푸시(/ws)"],
], widths=[18, 36, 112])
H2("4.3 티켓 라이프사이클")
P("관제팀은 접수·정/오탐 판정만, 대응 배정과 최종 종결은 정보보호 담당자가 결정한다. 우선순위별 SLA가 부여된다.")
TABLE(["상태", "담당", "다음 단계"], [
    ["미접수", "—", "접수(보안관제)"],
    ["접수", "보안관제팀", "검토(정탐 판정 후 정보보호 이관)"],
    ["검토", "정보보호팀", "대응(웹관리자/정보보호) 또는 오탐요청"],
    ["대응", "웹관리자/정보보호팀", "승인대기 또는 종결"],
    ["승인대기", "정보보호팀", "종결"],
    ["종결/오탐종결/무시종결", "정보보호팀", "(닫힘) — 종결코드 기록"],
], widths=[40, 60, 66])
TABLE(["우선순위", "P1", "P2", "P3", "P4"], [["SLA(기한)", "1시간", "4시간", "24시간", "72시간"]],
      header_fill=HDR_BLUE, widths=[30, 34, 34, 34, 34])

# ════════════════════════════════════════════════════════════════════
# 5. 데이터 처리 파이프라인
# ════════════════════════════════════════════════════════════════════
H1("5. 데이터 처리 파이프라인")
H2("5.1 상관탐지 수신")
P("Splunk 상관룰(CORR-0xx)이 매분 다중 장비 로그를 엮어 탐지 결과를 soc_notable_json 인덱스에 "
  "JSON으로 적재한다. 각 notable은 rule_id·rule_title·mitre·severity·risk_score와 함께 "
  "‘기여 이벤트(evidence)’ 배열(실제 트리거 로그)을 담는다. 분석플랫폼은 이를 수신해 매핑한다.")
CODE("index=soc_notable_json sourcetype=\"soc:notable:json\"\n"
     "| sort -_time | head 100 | fields _raw _time")
P("entity가 IP면 ‘공격 출발지(src_ip)’, 호스트명이면 ‘영향 자산(asset)’으로 구분 매핑한다. "
  "evidence의 페이로드(payload/injected_params)와 modsec 본문을 추출해 판정 근거로 활용한다.")
H2("5.2 AI 정·오탐 자동 판정")
P("탐지를 ‘유형+출발지+페이로드+시간’ 상관 그룹으로 묶고, 그룹마다 대표 1건만 LLM(gemini-2.5-flash, "
  "temperature 0)으로 판정한다. 신규 그룹만 호출하고 결과를 캐시해 비용을 절감한다. 페이로드형(WAF/IDS) "
  "탐지는 내용 기반, 상관·행위형 탐지는 근거(evidence)·요약 기반의 전용 프롬프트로 분기 판정한다.")
H2("5.3 위험도·MITRE 보정")
BULLET("결정적 공격구문(_strong_attack_signal): UNION SELECT·<script>·웹쉘 등 명백 구문이 있으면 AI 오탐을 정탐으로 보정")
BULLET("볼륨 보정(_flood_volume): DDoS 등 대량성은 단일 대상 집중 건수가 임계(기본 20) 이상이면 정탐 확정")
BULLET("상관탐지 가드레일: 룰이 다중 근거로 탐지한 건은 페이로드 부재로 자동 오탐 처리하지 않음(‘심각’ 등급은 사람 검토 강제)")
BULLET("MITRE: 상관룰이 큐레이션한 mitre를 권위값으로 사용, 없으면 공격유형 키워드로 추론(default_meta_for)")
H2("5.4 중복제거(Deduplication)")
P("상관탐지는 매분 새 notable을 만들고 근거가 매번 달라지므로, 휘발성 페이로드 지문을 키에서 제외하고 "
  "‘rule_id + 엔티티 + 일(日)’로 묶어 같은 인시던트가 하루 1건이 되게 한다. DDoS는 대상(dest) 기준으로 "
  "분산 출발지를 한 캠페인으로 병합한다.")
CODE("# 상관탐지 dedup 키 (한 인시던트 = 하루 1티켓)\n"
     "corr|{rule_id}|{src_ip 또는 asset}|{floor(epoch/86400)}\n"
     "# DDoS(대량성): 대상 기준 병합\n"
     "DDoS|dst:{dest_ip}|{30분 버킷}")
H2("5.5 정탐 전달 및 티켓팅")
P("정탐만 소통플랫폼 /api/events/ingest로 전달한다. 전송한 키를 forwarded_keys에 영구 기록해 같은 "
  "인시던트의 재전송을 막고, 실패 시 다음 주기에 재시도한다. 위험도는 AI 판정으로 상향 보정(하향 없음)되어 "
  "우선순위(P1~P4)·SLA가 부여된다.")

# ════════════════════════════════════════════════════════════════════
# 6. API 명세서
# ════════════════════════════════════════════════════════════════════
H1("6. API 명세서")
H2("6.1 분석플랫폼 API (:8800)")
TABLE(["Method", "Path", "설명"], [
    ["GET", "/health", "헬스 체크"],
    ["GET", "/api/crisis-level", "국내 사이버 위기경보단계(KISA/KrCERT)"],
    ["GET", "/api/dashboard/news", "보안 뉴스 헤드라인(RSS)"],
    ["GET", "/api/dashboard/advisories", "기관 보안권고·공지(KrCERT)"],
    ["GET", "/api/alerts", "최근 알림(상관 notable) — index, earliest, latest"],
    ["GET", "/api/alerts/summary", "금일 위협현황 집계(전체 건수/위험도/유형)"],
    ["GET", "/api/alerts/geo", "지구본 arc 집계(출발지→목적지 흐름) — limit"],
    ["GET", "/api/triage", "정·오탐 자동판정(최근) — earliest, limit"],
    ["GET", "/api/triage/history", "탐지 기록 전체 — earliest, head"],
    ["POST", "/api/triage", "명시 이벤트(또는 최근 알림) 자동 판정"],
    ["POST", "/api/forward", "정탐 즉시 전달(수동 트리거, 멱등)"],
    ["POST", "/api/search", "임의 SPL 검색 실행"],
    ["POST", "/api/analyze", "이벤트 Gemini 분석 + PDF 보고서 생성"],
    ["POST", "/api/chat", "로그 분석 챗봇(24h 로그 컨텍스트 Q&A)"],
    ["POST", "/webhook/splunk", "Splunk alert action 웹훅 수신"],
    ["GET", "/api/reports", "PDF 보고서 목록"],
    ["GET", "/api/reports/{filename}", "PDF 보고서 다운로드"],
    ["GET", "/api/intel/ip/{ip}", "단일 IP 위협 인텔리전스"],
    ["POST", "/api/intel/enrich", "이벤트 IP 일괄 TI 조회"],
    ["GET", "/api/cve/recent", "최근 고위험 CVE(CVSS≥7) — limit"],
    ["POST", "/api/cve/by-signatures", "시그니처-CVE 매핑"],
    ["GET", "/api/insights/trends", "공격 트렌드 통계 + LLM 해석 — days"],
    ["GET", "/api/insights/summary", "주간 위협 요약(경영진용)"],
], widths=[16, 60, 90], bsize=8.5)
H2("6.2 소통플랫폼 API (:8810)")
TABLE(["Method", "Path", "설명"], [
    ["POST", "/api/login", "로그인"],
    ["POST", "/api/signup", "회원가입(사서함 자동 연결)"],
    ["GET", "/api/users", "사용자 목록"],
    ["POST", "/api/events/ingest", "정탐 이벤트 수신(배치)"],
    ["POST", "/api/events/ticket", "수동 티켓 생성"],
    ["GET", "/api/events", "이벤트 목록(필터)"],
    ["GET", "/api/events/stats · /metrics · /progress", "통계·지표·진행현황"],
    ["GET", "/api/events/{id}", "이벤트 상세"],
    ["POST", "/api/events/{id}/status", "상태 변경"],
    ["POST", "/api/events/{id}/assign", "담당자 배정"],
    ["POST", "/api/events/{id}/comments", "댓글 등록"],
    ["POST", "/api/events/{id}/priority", "우선순위 변경"],
    ["PATCH", "/api/events/{id}/meta", "태그·MITRE 수정"],
    ["POST/DELETE", "/api/events/{id}/attachments[/{att}]", "증적 첨부 추가/삭제"],
    ["POST/PATCH/DELETE", "/api/events/{id}/tasks[/{task}]", "대응 작업 추가/토글/삭제"],
    ["GET", "/api/ledger · /api/ledger/export", "대장 조회 · CSV 내보내기"],
    ["GET/POST/PATCH/DELETE", "/api/iocs[...]", "IOC 목록/생성/수정/삭제·통계"],
    ["POST", "/api/iocs/extract/{event_id}", "이벤트에서 IOC 추출"],
    ["GET/POST", "/api/mail/external/inbox·sent·send", "받은/보낸편지함·발송(첨부 포함)"],
    ["POST/GET", "/api/mail/upload · external/attachment", "첨부 업로드 · 첨부 다운로드"],
    ["GET/POST", "/api/mail/external/trash·restore·purge", "휴지통 이동·복원·영구삭제"],
    ["GET/POST", "/api/mail/unread_count·external/read·account", "안읽음 수·읽음표시·사서함 연결상태"],
    ["GET", "/api/chat/channels", "채널 목록 조회"],
    ["GET/POST", "/api/chat/channels/{id}/messages", "채널 메시지 조회·전송"],
    ["GET/POST", "/api/dm[/threads·/conversation·/unread_count]", "DM 스레드·대화·발송·안읽음"],
    ["GET", "/api/departments · /api/contacts", "부서·연락처 주소록(추가/삭제)"],
    ["GET", "/api/mitre", "MITRE 기법 매핑 목록"],
    ["POST", "/api/upload", "파일 업로드(증적·이미지)"],
    ["WS", "/ws", "실시간 알림(이벤트·채팅·DM 푸시)"],
], widths=[26, 66, 74], bsize=8.5)

# ════════════════════════════════════════════════════════════════════
# 7. DB 명세서
# ════════════════════════════════════════════════════════════════════
H1("7. DB 명세서")
H2("7.1 소통플랫폼 데이터베이스 (SQLite)")
P("SQLAlchemy ORM 기반 12개 테이블. 시작 시 create_all + ALTER TABLE 방식의 경량 마이그레이션을 수행한다. "
  "핵심은 events(티켓)이며 comments·history·attachments·tasks가 1:N으로 연결된다.")

def DBT(title, rows):
    H3(title)
    TABLE(["컬럼", "타입", "설명"], rows, header_fill=HDR_BLUE, widths=[40, 34, 92], bsize=8.5)

DBT("users (사용자)", [
    ["id", "INT PK", "사용자 ID"],
    ["username", "VARCHAR(64) UQ", "로그인 ID"],
    ["display_name", "VARCHAR(64)", "표시 이름"],
    ["team", "VARCHAR(32)", "보안관제팀/웹관리자/정보보호팀"],
    ["role", "VARCHAR(32)", "역할(기본 ‘분석가’)"],
    ["password", "VARCHAR(128)", "로그인 비밀번호(데모 평문)"],
    ["email / phone", "VARCHAR", "알림 연락처"],
    ["notify_consent", "BOOL", "알림 수신 동의"],
    ["mail_address / mail_password", "VARCHAR", "janus.com 사서함 자격"],
    ["created_at", "DATETIME", "생성 시각"],
])
DBT("events (보안 이벤트·티켓)", [
    ["id", "INT PK", "이벤트 ID"],
    ["signature", "VARCHAR(255)", "탐지 시그니처/룰명"],
    ["src_ip / dest_ip", "VARCHAR(64)", "출발지/목적지 IP"],
    ["asset", "VARCHAR(128)", "호스트형 탐지 피해 자산"],
    ["uri / payload", "TEXT", "요청 경로 / 판정에 쓰인 공격 페이로드"],
    ["severity", "VARCHAR(8)", "위험도(1/2/3)"],
    ["source / detected_at", "VARCHAR", "탐지원 / 원본 시각"],
    ["ai_verdict / ai_confidence", "VARCHAR/INT", "AI 판정 / 신뢰도"],
    ["ai_attack_type / ai_reasoning", "VARCHAR/TEXT", "공격 유형 / 판정 근거"],
    ["dup_count", "INT", "중복 병합 건수"],
    ["ticket_no / priority / due_at", "VARCHAR/DATETIME", "티켓번호 / 우선순위 / SLA 기한"],
    ["tags / mitre", "VARCHAR", "태그 / MITRE 기법 ID(콤마)"],
    ["origin", "VARCHAR(16)", "분석플랫폼 / 수동"],
    ["status", "VARCHAR(16) IDX", "라이프사이클 상태"],
    ["assignee_id", "FK users", "담당자"],
    ["resolved_at / resolution_code / root_cause", "DATETIME/VARCHAR/TEXT", "종결 시각 / 종결코드 / 근본원인"],
    ["created_at / updated_at", "DATETIME", "생성 / 수정 시각"],
])
DBT("event_comments / event_history / event_attachments / event_tasks", [
    ["event_comments", "id, event_id FK, user_id FK, body TEXT, created_at", "티켓 댓글"],
    ["event_history", "id, event_id FK, user_id FK?, action, detail TEXT, created_at", "상태변경 감사 이력"],
    ["event_attachments", "id, event_id FK, url, name, size, uploaded_by_id FK?, created_at", "증적 첨부"],
    ["event_tasks", "id, event_id FK, title, done BOOL, created_at", "대응 작업/체크리스트"],
])
DBT("iocs (침해지표)", [
    ["id", "INT PK", "IOC ID"],
    ["ioc_type", "VARCHAR(16) IDX", "IP/도메인/URL/해시/이메일"],
    ["value", "VARCHAR(512) IDX", "지표 값"],
    ["severity / confidence", "VARCHAR/INT", "위험도 / 신뢰도"],
    ["status", "VARCHAR(16) IDX", "활성/차단완료/만료/오탐제외"],
    ["description", "TEXT", "설명"],
    ["source_event_id / created_by_id", "FK", "출처 이벤트 / 등록자"],
    ["first_seen / last_seen", "VARCHAR(40)", "최초/최근 관측"],
    ["created_at / updated_at", "DATETIME", "생성/수정"],
])
DBT("mails (사내 메일 메타)", [
    ["id", "INT PK", "메일 ID"],
    ["sender_id / recipient_id", "FK users", "발신/수신(인앱)"],
    ["recipient_email / name / dept", "VARCHAR", "외부 수신 정보"],
    ["channel", "VARCHAR(16)", "inapp / email"],
    ["subject / body", "VARCHAR/TEXT", "제목/본문"],
    ["is_read", "BOOL", "읽음 여부"],
    ["del_sender / del_recipient", "BOOL", "사용자별 소프트 삭제"],
    ["related_event_id", "FK events", "연관 이벤트"],
    ["created_at", "DATETIME", "생성 시각"],
])
DBT("channels / chat_messages / direct_messages / contacts", [
    ["channels", "id, name UQ, description, created_at", "채팅 채널"],
    ["chat_messages", "id, channel_id FK, user_id FK, body, event_id FK?, image_url?, created_at", "채널 메시지(이벤트카드·이미지)"],
    ["direct_messages", "id, sender_id FK, recipient_id FK, body, image_url?, is_read, created_at", "1:1 DM"],
    ["contacts", "id, name, email IDX, dept, note, created_at", "사내 부서 주소록"],
])
H2("7.2 분석플랫폼 데이터 소스 (무 RDB)")
P("분석플랫폼은 관계형 DB 대신 Splunk 인덱스를 데이터 소스로, 로컬 JSON 파일을 캐시로 사용한다.")
TABLE(["소스", "종류", "용도"], [
    ["soc_notable_json", "Splunk 인덱스", "상관룰 탐지(JSON, evidence 포함) — 1차 알림 소스"],
    ["soc_notable", "Splunk 인덱스", "구형 notable(key=value)"],
    ["soc_base 매크로", "Splunk 매크로", "WAF/IDS/pfSense 통합·정규화 원시 alert(집계·인사이트)"],
    ["triage_cache.json", "로컬 캐시", "dedup_key→판정결과(신규만 LLM 호출)"],
    ["forwarded_keys.json", "로컬 캐시", "전달 완료 키(재전송 방지)"],
    ["pdf_reports/", "로컬 저장", "생성된 PDF 보고서"],
], widths=[44, 32, 90])

# ════════════════════════════════════════════════════════════════════
# 8. AI · 상관분석 설계
# ════════════════════════════════════════════════════════════════════
H1("8. AI · 상관분석 설계")
H2("8.1 LLM 구성")
P("OpenRouter(OpenAI 호환)로 google/gemini-2.5-flash를 호출한다. 호출 구조는 system 프롬프트 + user "
  "메시지의 2-메시지 방식이며, 정·오탐 판정은 temperature 0(재현성)으로 고정한다. 용도: 정·오탐 판정, "
  "공격분석 보고서(요약 카드), 로그 분석 챗봇, 위협 인사이트.")
H2("8.2 판정 프롬프트(발췌)")
CODE("[판별 기준]\n"
     "- 정탐: 페이로드/근거에 실제 공격 의도가 명확(UNION SELECT, <script>, ../../etc/passwd ...)\n"
     "- 오탐: 시그니처는 매칭됐으나 정상 트래픽(정상 검색어/세션ID/인증된 스캐너 등)\n"
     "[원칙] 페이로드의 ‘실제 내용’을 근거로 판단, URL 인코딩은 디코딩해 해석,\n"
     "       애매하면 보안상 안전한 쪽(정탐 의심)으로 기운다. 결과는 JSON 배열로만 출력.")
H2("8.3 판정 신뢰성 설계")
BULLET("LLM 단독이 아니라 ‘결정적 룰 보정 + LLM 판정 + (상관) 근거 기반 판정’의 결합 파이프라인")
BULLET("내용형(payload) 탐지와 행위형(correlation) 탐지를 서로 다른 프롬프트·근거로 분기")
BULLET("Snort 단독 탐지처럼 본문이 없는 경우 ‘본문 미확보’로 정직하게 표기(가짜 페이로드 미생성)")

# ════════════════════════════════════════════════════════════════════
# 9. 보안 고려사항
# ════════════════════════════════════════════════════════════════════
H1("9. 보안 고려사항")
TABLE(["항목", "현황 / 권고"], [
    ["인증", "데모는 평문 비밀번호 — 운영 전 해시(bcrypt 등) 및 세션/토큰 적용 권고"],
    ["비밀값 관리", "Splunk·메일·LLM 자격은 .env 환경변수로만 주입(코드/문서 미포함)"],
    ["메일 게이트웨이", "DMZ 게이트웨이를 통한 송수신, 자체서명 인증서 환경(TLS 검증 옵션)"],
    ["메일 첨부", "실행 확장자 차단·파일당 10MB 상한·staging 웹 비공개(발송 후 삭제)·경로조작 방지"],
    ["외부 크롤링", "제목·링크만 수집·캐시·출처표기 — 본문 무단복제 금지(저작권/약관 준수)"],
    ["접근 범위", "내부망 운영 전제, 외부 인터넷은 FW 아웃바운드 룰로 제한적 허용"],
    ["감사 추적", "티켓 상태변경·배정·종결을 event_history로 기록"],
], widths=[34, 132])

# ════════════════════════════════════════════════════════════════════
# 10. 상관 탐지룰 및 MITRE 커버리지
# ════════════════════════════════════════════════════════════════════
H1("10. 상관 탐지룰 및 MITRE 커버리지")
H2("10.1 상관 탐지룰 목록")
P("Splunk 상관룰(soc_notable_json 적재)은 단일 장비 알림이 아니라 다중 근거를 엮어 ‘의미 있는 인시던트’를 "
  "탐지한다. 본 구축에서 설계·정의한 상관룰 카탈로그는 다음과 같다(일부 룰은 Splunk 서버 배포·고도화가 "
  "진행 중이다. 분류: 웹=네트워크 페이로드형, 호스트=엔드포인트 행위형, 대량성=볼륨/봇, 복합=다중룰 상관).")
TABLE(["rule_id", "탐지명", "MITRE", "분류"], [
    ["CORR-001", "SQLi 차단실패·성공 의심", "T1190", "웹"],
    ["CORR-010", "CSRF 상태변경(Referer 불일치)", "T1565.001 / T1185", "웹"],
    ["CORR-011", "정찰·스캐닝(포트스캔/탐침)", "T1595 / T1046", "웹"],
    ["CORR-012", "웹셸 업로드(위험확장자/시그)", "T1505.003", "웹"],
    ["CORR-013", "RCE·웹셸 명령실행", "T1059", "호스트"],
    ["CORR-014", "크리덴셜 브루트포스(로그인 rate)", "T1110", "대량성"],
    ["CORR-015", "데이터 유출 성공(ATTACK_RESPONSE)", "T1213 / T1041", "웹"],
    ["CORR-016", "악성 실행파일(ELF) 다운로드/전송", "T1105", "호스트"],
    ["CORR-017", "매크로/봇 자동화 티켓팅", "T1499.003", "대량성"],
    ["CORR-018", "복합 웹공격 킬체인(다단계)", "다단계", "복합"],
    ["CORR-019", "리버스셸·C2 에이전트 실행", "T1059 / T1571", "호스트"],
    ["CORR-020", "권한상승(privesc)", "T1548 / T1068", "호스트"],
    ["CORR-021", "지속성·백도어 설치", "T1053 / T1098 / T1136", "호스트"],
    ["CORR-022", "방어회피·로그변조(timestomp)", "T1070", "호스트"],
    ["CORR-023", "인가우회·권한파라미터 변조(Mass Assignment)", "T1068 / CWE-915", "웹"],
    ["CORR-024", "민감 설정·자격파일 노출", "T1552 / T1083", "웹"],
    ["CORR-025", "XSS 캠페인(다벡터·지속)", "T1059.007", "웹"],
    ["CORR-027", "호스트 침투 단계 상관(멀티키)", "다단계", "호스트"],
    ["CORR-028", "크리덴셜 스토어 접근(shadow)", "T1003", "호스트"],
    ["CORR-029", "스캘핑·매크로 봇(행위지표)", "T1499.003", "대량성"],
    ["CORR-030", "복합 침해(다중룰 동시발화)", "위험기반 상관", "복합"],
    ["CORR-031", "크리덴셜 스터핑·대량 계정탈취", "T1110.004 / T1078", "대량성"],
    ["CORR-007", "경로순회 파일 유출(LFI·Path Traversal)", "T1083 / T1005", "웹"],
    ["CORR-008", "저장형 XSS(게시판/프로필)", "T1059.007", "웹"],
    ["CORR-009", "XSS 세션쿠키 탈취 시도", "T1539 / T1059.007", "웹"],
    ["CORR-002 / 003", "다층 DDoS / L7 DoS(가용성)", "T1498 / T1499", "대량성"],
    ["CORR-005 / 006 ※", "랜섬웨어 킬체인 / 암호화", "T1486", "호스트"],
], widths=[34, 70, 44, 18], bsize=8.5)
NOTE("CORR-005/006(랜섬웨어)은 table형 collect 룰로 내부 rule_id가 RANSOM-KC-A / RANSOM-ENC-B이며, "
     "JSON notable(soc:notable:json)로 적재되지 않아 분석플랫폼 자동 수신 대상에서는 제외된다(별도 분석 트랙).",
     "표기 주석")
P("각 룰은 결과 JSON에 기여 이벤트(evidence)와 위험점수(risk_score)를 담아, 분석플랫폼이 근거 기반으로 "
  "정·오탐을 판정할 수 있게 한다.", size=10)
H2("10.2 MITRE ATT&CK 킬체인 커버리지")
TABLE(["전술(Tactic)", "기법(Technique)", "관련 룰"], [
    ["초기 침투(Initial Access)", "T1190 Exploit Public-Facing App", "CORR-001, 007"],
    ["실행(Execution)", "T1059, T1059.007", "CORR-013/019/025, 008"],
    ["지속성(Persistence)", "T1053, T1098, T1136", "CORR-021"],
    ["권한 상승(Priv. Escalation)", "T1548, T1068", "CORR-020/023"],
    ["방어 회피(Defense Evasion)", "T1070", "CORR-022"],
    ["자격 접근(Credential Access)", "T1003, T1110(.004), T1539", "CORR-014/028/031, 009"],
    ["발견/정찰(Discovery)", "T1046, T1083, T1595", "CORR-011/024, 007"],
    ["수집/유출(Exfiltration)", "T1005, T1213, T1041", "CORR-015/024, 007"],
    ["명령·제어(C2)", "T1105, T1571", "CORR-016/019"],
    ["영향(Impact)", "T1486, T1498, T1499(.003)", "CORR-005/006, DDoS, 봇"],
], widths=[44, 70, 52], bsize=8.5)

# ════════════════════════════════════════════════════════════════════
# 11. 인터페이스 정의서 (분석 → 소통)
# ════════════════════════════════════════════════════════════════════
H1("11. 인터페이스 정의서")
H2("11.1 분석플랫폼 → 소통플랫폼 연동")
P("정탐으로 판정된 이벤트는 분석플랫폼이 소통플랫폼의 수신 API로 배치 전송한다. 두 시스템 간 연동 계약은 "
  "아래 IngestEvent 스키마이다.")
TABLE(["구분", "내용"], [
    ["엔드포인트", "POST http://<소통플랫폼>:8810/api/events/ingest"],
    ["요청 형식", "JSON { \"events\": [ IngestEvent, ... ] }"],
    ["응답", "{ \"ingested\": N, \"ids\": [...] }"],
    ["트리거", "분석플랫폼 60초 백그라운드 루프(정탐 자동) / POST /api/forward(수동)"],
    ["멱등성", "dedup_key를 forwarded_keys에 영구 기록 → 재전송 방지"],
], widths=[28, 138])
H2("11.2 전달 페이로드(IngestEvent) 필드")
TABLE(["필드", "타입", "설명"], [
    ["signature", "str", "탐지 시그니처/룰명"],
    ["src_ip / dest_ip", "str", "출발지 / 목적지 IP"],
    ["asset", "str", "호스트형 탐지 피해 자산(호스트명)"],
    ["uri / payload", "str", "요청 경로 / 판정에 쓰인 실제 공격 페이로드"],
    ["severity", "str", "위험도(1/2/3) — AI 판정으로 상향 보정"],
    ["source", "str", "탐지원(SIEM 상관룰 등)"],
    ["mitre", "str", "MITRE 기법 ID(룰 큐레이션값)"],
    ["detected_at", "str", "원본 탐지 시각"],
    ["ai_verdict / ai_confidence", "str / int", "AI 판정(정탐) / 신뢰도(%)"],
    ["ai_attack_type / ai_reasoning", "str / str", "공격 유형 / 판정 근거"],
    ["dup_count", "int", "중복제거로 병합된 원본 경보 수"],
], widths=[40, 26, 100], bsize=8.5)
NOTE("위험도는 AI 판정·공격유형·신뢰도로 상향만 보정한다(하향 없음). 예: 규칙이 ‘WAF Anomaly(낮음)’로 "
     "분류한 웹쉘 업로드를 AI가 고위험으로 격상.", "위험도 보정")

# ════════════════════════════════════════════════════════════════════
# 12. 화면 정의서
# ════════════════════════════════════════════════════════════════════
H1("12. 화면 정의서")
H2("12.1 분석플랫폼 화면")
TABLE(["화면/탭", "구성", "기능"], [
    ["메인 대시보드", "금일 위협현황·위기경보 게이지·보안뉴스/권고 탭·지구본", "현황 요약 + 공격 출발지→목적지 호 시각화"],
    ["알림 탭", "최근 탐지 목록(중복 ×N 묶음)", "실시간 탐지 모니터링"],
    ["정·오탐 탭", "AI 판정 결과(근거·신뢰도)", "최근 자동 판정 검토"],
    ["기록 탭", "전체 데이터베이스 표(검색·필터·페이지)", "누적 탐지 이력 조회"],
    ["인텔리전스 탭", "IP 평판·CVE", "위협 인텔리전스 보강"],
    ["인사이트 탭", "트렌드 통계 + AI 해석", "기간 분석·주간 요약"],
    ["보고서 탭", "PDF 보고서 목록·생성", "공격 분석 보고서"],
    ["보안 분석 어시스턴트", "플로팅 챗봇", "로그 자연어 Q&A"],
], widths=[30, 66, 70], bsize=8.5)
H2("12.2 소통플랫폼 화면")
TABLE(["화면", "기능"], [
    ["대시보드", "처리 지표·진행현황·최근 이벤트"],
    ["이벤트(티켓)", "목록·상세(상태/배정/댓글/작업/첨부/페이로드/MITRE)"],
    ["대장(원장)", "전체 처리 대장 표·CSV 내보내기"],
    ["IOC", "침해지표 등록·차단상태 관리·이벤트 추출"],
    ["메일", "받은/보낸/휴지통·검색·작성·첨부파일 송수신(janus.com)"],
    ["채팅 / DM", "팀 채널 채팅 · 1:1 다이렉트 메시지"],
    ["연락처", "사내 부서 주소록"],
], widths=[34, 132])

# ════════════════════════════════════════════════════════════════════
# 13. 테스트 및 검증 결과
# ════════════════════════════════════════════════════════════════════
H1("13. 테스트 및 검증 결과")
H2("13.1 AI 정·오탐 판정")
TABLE(["항목", "결과"], [
    ["명백 공격 판정", "SQLi·XSS·웹쉘 등 결정적 구문 보유 건 정탐 일치"],
    ["거짓음성 보정", "LLM이 ‘페이로드 없음’으로 오판한 저장형 XSS(<svg onload=alert(1)>)를 결정적 보정으로 정탐 교정"],
    ["행위형 판정", "세션 하이재킹(페이로드 없음)을 근거(요약) 기반으로 정탐 판정(신뢰도 80%)"],
    ["가드레일", "상관탐지는 페이로드 부재만으로 자동 오탐 처리하지 않음 / ‘심각’ 등급은 사람 검토 강제"],
], widths=[34, 132])
H2("13.2 중복제거 및 전달")
TABLE(["항목", "전 → 후"], [
    ["WAF Anomaly 노이즈(soc_base 필터)", "772 → 67건"],
    ["호스트룰 중복 유입", "약 78초마다 신규 → 하루 1티켓(인시던트당)"],
    ["누적 중복 티켓 정리", "356 → 98건(작업본 보존, 미접수 중복만 삭제)"],
    ["DDoS 분산 출발지", "수천 출발지 → 대상 기준 1캠페인(1티켓)"],
], widths=[60, 106])
H2("13.3 페이로드·자산 추출 검증")
BULLET("저장형 XSS: 본문에서 `<svg onload=alert(1)>` 정확 추출, 소통 티켓 ‘판정 페이로드’에 표시")
BULLET("Mass Assignment: 앱로그에서 `role=ADMIN, MEMBER_GRADE=ADMIN` 추출(IDS 단독 건은 ‘본문 미확보’로 정직 표기)")
BULLET("호스트형 탐지: entity가 호스트명이면 ‘영향 자산(🖥)’으로 분리 표시(출발지 오표기 교정)")
BULLET("메일 첨부 송수신: 발송→수신→다운로드 E2E 검증 — 한글 파일명 보존·바이트 일치(PASS)")
H2("13.4 비용 최적화")
P("동일 인시던트의 대표 1건만, 그것도 신규(미캐시) 그룹만 LLM 호출하고 결과를 캐시·배치 처리하여 LLM "
  "토큰 사용량을 대폭 절감했다(중복제거+배치 적용 전 대비 약 98% 절감).")

# ════════════════════════════════════════════════════════════════════
# 14. 환경 구성 및 구축 이력
# ════════════════════════════════════════════════════════════════════
H1("14. 환경 구성 및 구축 이력")
H2("14.1 환경변수(.env)")
TABLE(["변수", "용도"], [
    ["SPLUNK_HOST / PORT / TOKEN / USERNAME / PASSWORD", "Splunk REST 접속(토큰 우선)"],
    ["GEMINI_API_KEY / GEMINI_MODEL", "OpenRouter LLM(기본 gemini-2.5-flash)"],
    ["COMM_PLATFORM_URL / COMM_FORWARD_ENABLED", "소통플랫폼 전달 대상·on/off"],
    ["SOC_ALERT_SOURCE", "1차 알림 소스(notable | socbase)"],
    ["TRIAGE_FLOOD_MIN / TRIAGE_SPLIT_BY_PAYLOAD", "볼륨 보정 임계 / 페이로드별 분리"],
    ["MAIL_GATEWAY_HOST / MAIL_DOMAIN / 포트", "메일 게이트웨이(janus.com, IMAP993/SMTP25)"],
], widths=[64, 102], bsize=8.5)
H2("14.2 Splunk 구성")
BULLET("soc_base 매크로: WAF Web(modsec) + IDS/Snort + pfSense를 통합·정규화한 원시 alert 베이스")
BULLET("상관룰(CORR-0xx): 매분 스케줄 검색 → soc_notable_json(JSON, evidence 포함) 적재")
BULLET("‘SOC 통합 관제(SIEM)’ 대시보드: 상관탐지 현황 패널")
H2("14.3 구축 중 주요 이슈 및 개선")
TABLE(["이슈", "원인", "개선"], [
    ["탐지 이벤트 100 고정", "알림 쿼리 head 100 제한", "전체 집계 엔드포인트 분리(/summary)"],
    ["DDoS·행위형 오탐", "내용 기반 LLM이 무페이로드를 오탐", "볼륨/상관 가드레일·전용 프롬프트 도입"],
    ["같은 알림 반복 전달", "휘발성 페이로드 지문이 dedup 키를 흔듦 + 다중 프로세스", "키를 rule_id+엔티티+일로 안정화, 단일 인스턴스화"],
    ["페이로드 일부만 표시", "uri만 전달, 필드명 불일치(url/injected_params)", "실제 페이로드 추출·전달, 대체 필드 매핑"],
    ["출발지에 호스트명 표기", "호스트형 entity를 src_ip로 매핑", "IP=출발지 / 호스트명=영향 자산 분리"],
    ["MITRE 부정확", "키워드 추론이 한글 유형 미매칭→T1190", "룰 큐레이션 mitre 우선 사용"],
], widths=[40, 60, 66], bsize=8.5)
H2("14.4 변경 이력")
TABLE(["버전", "일자", "내용"], [
    ["1.0", TODAY, "최초 작성 — 구축 결과 정리(기능·API·DB·상관룰·테스트)"],
    ["1.1", TODAY, "메일 첨부파일 송수신 기능 추가 반영(업로드/다운로드 API·화면·검증)"],
    ["1.2", REVDATE, "전 시스템 정합성 점검 — 채팅 라우트 표기, 상관룰 ID(007~009)·랜섬웨어 룰 각주 정정"],
    ["1.3", REVDATE, "문서 서식을 WAS 구축보고서 템플릿에 맞춰 재구성(간결 표지·■ 목차·인라인 캡션/주석·머리글 제거, 페이지번호 유지)"],
], widths=[20, 32, 114])

# ════════════════════════════════════════════════════════════════════
# 15. 운영 및 배포
# ════════════════════════════════════════════════════════════════════
H1("15. 운영 및 배포")
H2("15.1 실행")
CODE("# 분석플랫폼\ncd 분석플랫폼/backend\n.venv\\Scripts\\python -m uvicorn main:app --host 0.0.0.0 --port 8800\n\n"
     "# 소통플랫폼\ncd 소통플랫폼/backend\n.venv\\Scripts\\python -m uvicorn main:app --host 0.0.0.0 --port 8810")
H2("15.2 운영 유의사항")
BULLET("백엔드는 포트별 단일 인스턴스로 기동(중복 기동 시 전달 루프 중복·티켓 중복 발생)")
BULLET("정탐 전달은 60초 주기 백그라운드 루프 + forwarded_keys로 멱등 보장")
BULLET("프런트엔드는 빌드 산출물(build/dist)을 백엔드가 정적 서빙 — 변경 시 재빌드")
BULLET("향후 우분투 VM(단일 호스트)에 두 플랫폼을 SSH로 배포 예정")

# ════════════════════════════════════════════════════════════════════
# 11. 결론 및 향후 과제
# ════════════════════════════════════════════════════════════════════
H1("16. 결론 및 향후 과제")
P("Splunk 상관분석 → AI 정·오탐 판정 → 협업 티켓팅으로 이어지는 end-to-end 보안관제 체계를 구축하여, "
  "대량 단일 알림을 ‘의미 있는 인시던트’ 중심으로 처리하고 팀 간 협업을 표준화했다.")
H2("향후 과제")
BULLET("인증 강화(비밀번호 해시·권한 분리·세션)")
BULLET("상관룰 확대 및 RBA(위험기반 알림) 도입, notable 기여 이벤트에 WAF 본문 enrichment")
BULLET("LLM 티어드 라우팅(쉬운 건 flash, 애매·고위험 건 상위 모델)·평가셋 기반 정확도 측정")
BULLET("우분투 VM 배포 자동화 및 모니터링")

# updateFields는 넣지 않는다(넣으면 '열 때마다' 갱신 프롬프트가 뜸).
# 대신 생성 후 Word 자동화(bake_fields.ps1)로 목차·페이지번호를 1회 계산해 구워넣는다.
# → 사용자는 열 때 프롬프트 없이 채워진 목차를 보고, F9도 누를 필요 없음.

doc.save(str(OUT))
print("OK ->", OUT)
print("문단/표 작성 완료 · 필드는 생성 후 Word 자동화로 bake")
